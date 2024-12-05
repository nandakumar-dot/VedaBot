import os
import base64
import io
import sys
from typing import Optional, List, Dict
from dotenv import load_dotenv
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.output_parsers import StrOutputParser
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from groq import Groq
from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
from database import Database

class TelegramAIAssistant:
    def __init__(self, telegram_token: str, google_api_key: str, groq_api_key: str, mongodb_uri: str):
        self.telegram_token = telegram_token        
        try:
            from google.generativeai import configure
            configure(api_key=google_api_key)
            self.groq_client = Groq(api_key=groq_api_key)
        except Exception as e:
            print(f"API Configuration Error: {e}")
            sys.exit(1)

        try:
            self.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        except Exception as e:
            print(f"Vector Store Initialization Error: {e}")
            sys.exit(1)

        self.db = Database(mongodb_uri)

        self.qa_chain = self._get_qa_chain()

    def _get_qa_chain(self):
        prompt_template = """
        You are a helpful AI assistant answering questions using provided context and general knowledge.

        Context:\n {context}\n
        Question:\n {question}\n

        Answer:
        """
        
        model = ChatGoogleGenerativeAI(
            model="gemini-pro", 
            temperature=0.3, 
            top_p=0.8, 
            top_k=40, 
            max_output_tokens=2048
        )
        
        prompt = PromptTemplate(
            template=prompt_template, 
            input_variables=["context", "question"]
        )
        
        def format_docs(docs: List[Document]) -> str:
            return "\n\n".join(doc.page_content for doc in docs)
        
        return (
            {"context": lambda x: format_docs(x['docs']), "question": lambda x: x['question']}
            | prompt
            | model
            | StrOutputParser()
        )

    def extract_text_from_file(self, file_bytes: bytes, file_name: str) -> Optional[str]:

        file_like = io.BytesIO(file_bytes)
        
        try:
            if file_name.lower().endswith('.pdf'):
                reader = PdfReader(file_like)
                return " ".join(page.extract_text() for page in reader.pages)
            
            elif file_name.lower().endswith('.docx'):
                doc = DocxDocument(file_like)
                return "\n".join(paragraph.text for paragraph in doc.paragraphs)
            
            return None
        except Exception as e:
            print(f"Text extraction error for {file_name}: {e}")
            return None

    async def process_image_with_groq(self, file_bytes: bytes) -> Optional[str]:
        try:
            base64_image = base64.b64encode(file_bytes).decode('utf-8')
            completion = self.groq_client.chat.completions.create(
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Extract text from this image:"},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ]
                }],
                model="llama-3.2-11b-vision-preview",
            )
            return completion.choices[0].message.content
        except Exception as e:
            print(f"Image processing error: {e}")
            return None

    async def handle_message(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        user_id = str(update.message.from_user.id)
        user_name = update.message.from_user.full_name
        self.db.add_user(user_id, user_name)
        try:
            user_question = update.message.text

            user_embeddings = self.db.get_user_embeddings(user_id)
            if not user_embeddings:
                await update.message.reply_text("No context available. Please upload documents first.")
                return
            
            docs = [
                Document(page_content=e["metadata"]["context"]) 
                for e in user_embeddings
            ]
            response_text = self.qa_chain.invoke({'docs': docs, 'question': user_question})
            
            await update.message.reply_text(response_text)

        except Exception as e:
            print(f"Message handling error: {e}")
            await update.message.reply_text("Error: Unable to process your request.")


    async def handle_file(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        user_id = str(update.message.from_user.id)
        user_name = update.message.from_user.full_name
        self.db.add_user(user_id, user_name)

        try:
            file = await update.message.document.get_file()
            file_bytes = await file.download_as_bytearray()
            file_name = update.message.document.file_name
            mime_type = update.message.document.mime_type

            text = None
            if mime_type == 'application/pdf' or file_name.lower().endswith('.pdf'):
                text = self.extract_text_from_file(file_bytes, file_name)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_name.lower().endswith('.docx'):
                text = self.extract_text_from_file(file_bytes, file_name)
            elif mime_type.startswith('image/'):
                text = await self.process_image_with_groq(file_bytes)

            if text:
                document_id = self.db.add_document(user_id, file_name, mime_type, text)
                
                embedding_vector = self.embeddings.embed_documents([text])
                self.db.add_embedding(user_id, document_id, embedding_vector[0], context=text[:200])

                await update.message.reply_text(f"✅ {file_name} processed successfully!")
            else:
                await update.message.reply_text("⚠️ Unable to extract text from the file.")
        
        except Exception as e:
            print(f"File handling error: {e}")
            await update.message.reply_text("⚠️ Error processing the file.")


    async def start_command(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        welcome_message = (
            "👋 Welcome to your AI Assistant!\n\n"
            "I can:\n"
            "1. 📚 Learn from uploaded documents (PDF, DOCX).\n"
            "2. 🖼️ Extract text from images.\n"
            "3. ❓ Answer questions using context or general knowledge.\n\n"
            "Upload a file or ask a question to get started!"
        )
        await update.message.reply_text(welcome_message)

    def run(self):
        load_dotenv()

        application = Application.builder().token(self.telegram_token).build()
        
        application.add_handler(CommandHandler("start", self.start_command))
        application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, self.handle_message))
        application.add_handler(MessageHandler(filters.Document.ALL, self.handle_file))
        
        application.run_polling(drop_pending_updates=True)

def main():

    assistant = TelegramAIAssistant(
        telegram_token=os.getenv("TELEGRAM_TOKEN"),
        google_api_key=os.getenv("GOOGLE_API_KEY"),
        groq_api_key=os.getenv("GROQ_API_KEY"),
        mongodb_uri=os.getenv("MONGODB_URI")
    )
    
    assistant.run()

if __name__ == '__main__':
    main()